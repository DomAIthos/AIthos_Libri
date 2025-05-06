from docx import Document

def crea_libro_docx(titolo, capitoli):
    doc = Document()
    doc.add_heading(titolo, 0)

    for i, (titolo_cap, testo) in enumerate(capitoli):
        doc.add_heading(f"Capitolo {i+1}: {titolo_cap}", level=1)
        doc.add_paragraph(testo)
        doc.add_page_break()

    nome_file = f"{titolo.replace(' ', '_')}.docx"
    doc.save(nome_file)
    print(f"Libro salvato come: {nome_file}")

# ESEMPIO DI USO
if __name__ == "__main__":
    titolo = "Il Viaggio della Mente"
    capitoli = [
        ("L'Inizio", "C'era una volta una mente curiosa..."),
        ("La Scoperta", "Un giorno, qualcosa cambiò..."),
    ]
    crea_libro_docx(titolo, capitoli)
